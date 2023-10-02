#from lms_tools.docx import DocxGenerator
from lms_tools.gift import GiftQuiz, GiftQuestion, GiftDistractor
from lms_tools.gift._utils import _escape_special_chars


def test_gift_quiz_docx_generator():
    quiz = GiftQuiz()
    q = GiftQuestion("L'appareil servant à mesurer la vitesse du vent au sol s'appelle :", name="0001", comment="question: 1 name: 0001")
    q.append_distractor(GiftDistractor("une girouette.", 0))
    q.append_distractor(GiftDistractor("une rose des vents.", 0))
    q.append_distractor(GiftDistractor("un baromètre.", 0))
    q.append_distractor(GiftDistractor("un anémomètre.", 1))
    quiz.append(q)

    q = GiftQuestion("L'unité de pression utilisée dans le système international et en aéronautique est :", name="0002", comment="question: 2 name: 0002")
    q.append_distractor(GiftDistractor("le pascal.", 1))
    q.append_distractor(GiftDistractor("le newton.", 0))
    q.append_distractor(GiftDistractor("le joule.", 0))
    q.append_distractor(GiftDistractor("le millimètre de mercure.", 0))
    quiz.append(q)

    #generator = DocxGenerator()
    from docx import Document
    document = Document()

    show_question_name = False
    show_correct_answers = True

    title = 'GIFT Quiz'
    if show_correct_answers:
        title += " (correction)"

    fname_out = 'gift_quiz.docx'

    document.add_heading(title, 0)

    p = document.add_paragraph('This is a ')
    p.add_run('GIFT').bold = True
    p.add_run(' quiz with ')
    p.add_run('great').italic = True
    p.add_run(' questions.')

    for i, q in enumerate(quiz.iter_questions()):
        heading = "Question %d" % (i + 1)
        if q.name != "" and show_question_name:
            heading += " (%s)" % q.name
        document.add_heading(heading, level=1)
        document.add_paragraph(q.stem)
        for j, distractor in enumerate(q.iter_distractors()):
            if show_correct_answers and q.is_partially_correct_answer(j) > 0:
                checkbox = "☑"
            else:
                checkbox = "☐"
            document.add_paragraph("\t" + checkbox + " " * 3 + _escape_special_chars(distractor.text))

    # document.add_page_break()

    document.save(fname_out)